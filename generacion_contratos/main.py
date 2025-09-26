import os
import subprocess

import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml import ns

import constants as const


def reemplazar_texto(
    parrafo, reemplazos: dict, fuente: str, tamaño_fuente: int
) -> None:
    """
    Reemplaza texto en un párrafo de Word, aplicando fuente y tamaño.
    """
    for buscar, reemplazo in reemplazos.items():
        if buscar in parrafo.text:
            nuevo_texto = parrafo.text.replace(buscar, reemplazo)
            parrafo.clear()
            ejecucion = parrafo.add_run(nuevo_texto)
            ejecucion.font.size = Pt(tamaño_fuente)
            ejecucion.font.name = fuente
            rFonts = ejecucion._element.rPr.rFonts
            rFonts.set(ns.qn("w:ascii"), fuente)
            rFonts.set(ns.qn("w:hAnsi"), fuente)
            rFonts.set(ns.qn("w:eastAsia"), fuente)
            rFonts.set(ns.qn("w:cs"), fuente)


def docx_a_pdf(ruta_docx: str, carpeta_salida: str) -> None:
    """
    Convierte un archivo .docx a PDF usando LibreOffice.
    """
    ruta_libreoffice = r"C:\\Program Files\\LibreOffice\\program\\soffice.exe"

    if not os.path.exists(ruta_libreoffice):
        raise FileNotFoundError("No se encontró LibreOffice.")

    comando = [
        ruta_libreoffice,
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        carpeta_salida,
        ruta_docx,
    ]
    subprocess.run(comando, check=True)


def convertir_todos_docx_a_pdf(carpeta_salida: str) -> None:
    """
    Convierte todos los archivos .docx de la carpeta de salida a PDF.
    """
    for archivo in os.listdir(carpeta_salida):
        if archivo.endswith(".docx"):
            nombre_base: str = archivo.replace(".docx", "")
            ruta_docx: str = os.path.join(carpeta_salida, archivo)

            try:
                docx_a_pdf(ruta_docx, carpeta_salida)
                print(f"Documento convertido {nombre_base}.pdf ✅")
            except Exception as e:
                print(f"Error al convertir {archivo} a PDF: {e}")


def generar_documentos(
    ruta_excel: str, ruta_plantilla: str, carpeta_salida: str
) -> None:
    """
    Genera documentos Word a partir de un Excel y una plantilla,
    y luego los convierte a PDF.
    """
    df: pd.DataFrame = pd.read_excel(ruta_excel)
    os.makedirs(carpeta_salida, exist_ok=True)
    print("\nIniciando generación de Word...\n")

    for _, fila in df.iterrows():
        doc = Document(ruta_plantilla)
        reemplazos = {
            clave: fila.get(valor, "") for clave, valor in const.MARCADORES.items()
        }

        for parrafo in doc.paragraphs:
            reemplazar_texto(
                parrafo, reemplazos, const.TIPO_FUENTE, const.TAMANO_FUENTE
            )

        nombre_archivo: str = f"{fila[const.CAMPO_SALIDA]}.docx"
        ruta_docx: str = os.path.join(carpeta_salida, nombre_archivo)
        doc.save(ruta_docx)

        print(f"Documento creado {nombre_archivo} ✅")

    print("\nIniciando conversión a PDF...\n")
    convertir_todos_docx_a_pdf(carpeta_salida)


if __name__ == "__main__":
    generar_documentos(
        const.ARCHIVO_ENTRADA, const.PLANTILLA_DOCX, const.CARPETA_SALIDA
    )
